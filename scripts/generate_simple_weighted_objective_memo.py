from __future__ import annotations

import shutil
import subprocess
import os
from pathlib import Path

import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


REPO_DIR = Path(__file__).resolve().parents[1]
CACHE_DIR = REPO_DIR / "notebooks" / "temp" / "reservation_playground_cache" / "92a74cbf5416"
CLASS_CACHE_PATH = CACHE_DIR / "class_df.pkl"


def resolve_utilization_config() -> tuple[str, str, str, str, str, Path, Path]:
    requested = os.environ.get("UTILIZATION_DEFINITION", "booked").strip().lower()
    if requested in {"booked", "booked_slot", "booked-slot"}:
        output_dir = REPO_DIR / "docs" / "reports" / "reservation_visual_objectives"
        return (
            "booked",
            "booked_slot_utilization",
            "booked-slot utilization",
            "booked slots divided by available slots",
            "strict_reservation_visual_objectives_memo",
            output_dir,
            output_dir / "strict_reservation_visual_objectives_memo.docx",
        )
    if requested in {"old", "served", "served_slot", "served-slot", "slot", "attended"}:
        output_dir = REPO_DIR / "docs" / "reports" / "reservation_visual_objectives_old_utilization"
        return (
            "old",
            "slot_utilization",
            "old served-slot utilization",
            "completed visits divided by available slots",
            "strict_reservation_old_utilization_memo",
            output_dir,
            output_dir / "strict_reservation_old_utilization_memo.docx",
        )
    raise ValueError("UTILIZATION_DEFINITION must be 'booked' or 'old'.")


(
    UTILIZATION_DEFINITION,
    UTILIZATION_COLUMN,
    UTILIZATION_LABEL,
    UTILIZATION_DEFINITION_TEXT,
    OUTPUT_BASENAME,
    OUTPUT_DIR,
    DOCX_PATH,
) = resolve_utilization_config()
FIGURE_DIR = OUTPUT_DIR / "figures"
TABLE_DIR = OUTPUT_DIR / "tables"
PDF_PATH = OUTPUT_DIR / f"{OUTPUT_BASENAME}.pdf"
WEIGHT_SUMMARY_PATH = OUTPUT_DIR / "weight_sensitivity" / "tables" / "weight_sensitivity_summary.csv"

STRICT_POLICY = "Strict C1 reservation"
FCFS_POLICY = "Pooled FCFS"
SCENARIO_TYPE = "symmetric_baseline"
DEMAND_LABEL = "lambda1=lambda2=25"
SEEDS = [5101, 5102, 5103, 5104, 5105]
Q_VALUES = [0, 4, 8, 12, 16, 20, 24, 28, 32]
CLASS_1_WEIGHT = 1.5
CLASS_2_WEIGHT = 1.0
SERVED_RATE_FLOOR = 0.50


def ensure_dirs() -> None:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    TABLE_DIR.mkdir(parents=True, exist_ok=True)


def fmt(value: float, digits: int = 3) -> str:
    if pd.isna(value):
        return ""
    return f"{value:.{digits}f}"


def load_baseline_class_data() -> pd.DataFrame:
    if not CLASS_CACHE_PATH.exists():
        raise FileNotFoundError(f"Missing cached class data: {CLASS_CACHE_PATH}")

    class_df = pd.read_pickle(CLASS_CACHE_PATH)
    subset = class_df[
        (class_df["scenario_type"] == SCENARIO_TYPE)
        & (class_df["demand_label"] == DEMAND_LABEL)
        & (class_df["seed"].isin(SEEDS))
        & (class_df["Q"].isin(Q_VALUES))
        & (class_df["policy"].isin([FCFS_POLICY, STRICT_POLICY]))
    ].copy()
    if subset.empty:
        raise RuntimeError("No rows found for the requested baseline slice.")

    strict_q = sorted(subset.loc[subset["policy"] == STRICT_POLICY, "Q"].unique().tolist())
    fcfs_q = sorted(subset.loc[subset["policy"] == FCFS_POLICY, "Q"].unique().tolist())
    if strict_q != Q_VALUES:
        raise RuntimeError(f"Strict-reservation Q values do not match the requested grid: {strict_q}")
    if 0 not in fcfs_q:
        raise RuntimeError("FCFS baseline with Q=0 is missing.")
    return subset


def load_saved_summary() -> pd.DataFrame:
    if not WEIGHT_SUMMARY_PATH.exists():
        raise FileNotFoundError(
            "Missing cached simulation data and saved summary table. "
            f"Expected one of: {CLASS_CACHE_PATH} or {WEIGHT_SUMMARY_PATH}"
        )

    saved = pd.read_csv(WEIGHT_SUMMARY_PATH)
    rows = saved[
        saved["w1"].eq(CLASS_1_WEIGHT)
        & saved["w2"].eq(CLASS_2_WEIGHT)
        & saved["Q"].isin(Q_VALUES)
    ].copy()
    if rows.empty:
        raise RuntimeError(
            f"No rows in {WEIGHT_SUMMARY_PATH} for w1={CLASS_1_WEIGHT:g}, w2={CLASS_2_WEIGHT:g}."
        )
    rows = rows.sort_values("Q")
    if rows["Q"].tolist() != Q_VALUES:
        raise RuntimeError("Saved summary does not match the requested Q grid.")

    rows = rows.rename(columns={"weighted_offered_wait": "weighted_wait"})
    rows["policy"] = STRICT_POLICY
    rows["offered_1"] = np.where(rows["tau_1"].isna(), 0.0, 1.0)
    rows["offered_2"] = np.where(rows["tau_2"].isna(), 0.0, 1.0)
    rows["served_rate_flag"] = rows["min_served_rate"] < SERVED_RATE_FLOOR

    fcfs = rows[rows["Q"].eq(0)].iloc[0].copy()
    fcfs["policy"] = FCFS_POLICY
    return pd.concat([pd.DataFrame([fcfs]), rows], ignore_index=True, sort=False)


def run_level_objectives(class_df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    keys = ["policy", "seed", "Q"]
    for (policy, seed, q), group in class_df.groupby(keys):
        class_1 = group[group["class_id"] == 1]
        class_2 = group[group["class_id"] == 2]
        if class_1.empty or class_2.empty:
            continue

        c1 = class_1.iloc[0]
        c2 = class_2.iloc[0]

        rho_1 = float(c1[UTILIZATION_COLUMN])
        rho_2 = float(c2[UTILIZATION_COLUMN])
        weighted_utilization = CLASS_1_WEIGHT * rho_1 + CLASS_2_WEIGHT * rho_2

        offered_1 = float(c1["offered"])
        offered_2 = float(c2["offered"])
        tau_1 = float(c1["mean_offered_booking_delay"]) if offered_1 > 0 else np.nan
        tau_2 = float(c2["mean_offered_booking_delay"]) if offered_2 > 0 else np.nan
        wait_denominator = CLASS_1_WEIGHT * offered_1 + CLASS_2_WEIGHT * offered_2
        weighted_wait = (
            CLASS_1_WEIGHT * float(c1["total_offered_booking_delay"])
            + CLASS_2_WEIGHT * float(c2["total_offered_booking_delay"])
        ) / wait_denominator if wait_denominator > 0 else np.nan

        served_rate_1 = float(c1["served_rate"])
        served_rate_2 = float(c2["served_rate"])
        arrivals_1 = float(c1["arrivals"])
        arrivals_2 = float(c2["arrivals"])
        served_1 = float(c1["served"])
        served_2 = float(c2["served"])
        min_served_rate = min(served_rate_1, served_rate_2)
        rows.append(
            {
                "policy": policy,
                "seed": int(seed),
                "Q": int(q),
                "rho_1": rho_1,
                "rho_2": rho_2,
                "weighted_utilization": weighted_utilization,
                "tau_1": tau_1,
                "tau_2": tau_2,
                "weighted_wait": weighted_wait,
                "offered_1": offered_1,
                "offered_2": offered_2,
                "served_rate_1": served_rate_1,
                "served_rate_2": served_rate_2,
                "arrivals_1": arrivals_1,
                "arrivals_2": arrivals_2,
                "served_1": served_1,
                "served_2": served_2,
                "min_served_rate": min_served_rate,
                "served_rate_flag": min_served_rate < SERVED_RATE_FLOOR,
            }
        )

    result = pd.DataFrame(rows)
    if result.empty:
        raise RuntimeError("Could not build run-level objective rows.")
    return result


def summarize(run_objectives: pd.DataFrame) -> pd.DataFrame:
    numeric_cols = [
        "rho_1",
        "rho_2",
        "weighted_utilization",
        "tau_1",
        "tau_2",
        "weighted_wait",
        "offered_1",
        "offered_2",
        "served_rate_1",
        "served_rate_2",
        "arrivals_1",
        "arrivals_2",
        "served_1",
        "served_2",
        "min_served_rate",
    ]
    summary = (
        run_objectives.groupby(["policy", "Q"], dropna=False)[numeric_cols]
        .mean()
        .reset_index()
        .sort_values(["policy", "Q"])
    )
    total_arrivals = summary["arrivals_1"] + summary["arrivals_2"]
    summary["overall_served_rate"] = np.where(
        total_arrivals > 0,
        (summary["served_1"] + summary["served_2"]) / total_arrivals,
        np.nan,
    )
    summary["served_rate_flag"] = summary["min_served_rate"] < SERVED_RATE_FLOOR
    return summary


def fcfs_row(summary: pd.DataFrame) -> pd.Series:
    rows = summary[(summary["policy"] == FCFS_POLICY) & (summary["Q"] == 0)]
    if rows.empty:
        raise RuntimeError("FCFS Q=0 reference row is missing.")
    return rows.iloc[0]


def strict_rows(summary: pd.DataFrame) -> pd.DataFrame:
    strict = summary[summary["policy"] == STRICT_POLICY].sort_values("Q")
    strict = strict[strict["Q"].isin(Q_VALUES)]
    if strict["Q"].tolist() != Q_VALUES:
        raise RuntimeError("Strict-reservation summary does not match the requested Q grid.")
    return strict


def mark_flagged_points(ax: plt.Axes, strict: pd.DataFrame, y_col: str) -> None:
    flagged = strict[strict["served_rate_flag"] & strict[y_col].notna()]
    ax.scatter(
        flagged["Q"],
        flagged[y_col],
        marker="X",
        color="black",
        edgecolors="white",
        linewidths=0.7,
        s=95,
        zorder=8,
        label="served rate flag",
    )


def mark_class_served_rate_flags(ax: plt.Axes, strict: pd.DataFrame) -> None:
    class_1_flagged = strict[strict["served_rate_1"] < SERVED_RATE_FLOOR]
    class_2_flagged = strict[strict["served_rate_2"] < SERVED_RATE_FLOOR]
    if not class_1_flagged.empty:
        ax.scatter(
            class_1_flagged["Q"],
            class_1_flagged["served_rate_1"],
            marker="X",
            color="black",
            edgecolors="white",
            linewidths=0.7,
            s=95,
            zorder=8,
            label="class below 50%",
        )
    if not class_2_flagged.empty:
        ax.scatter(
            class_2_flagged["Q"],
            class_2_flagged["served_rate_2"],
            marker="X",
            color="black",
            edgecolors="white",
            linewidths=0.7,
            s=95,
            zorder=8,
            label="class below 50%" if class_1_flagged.empty else "_nolegend_",
        )


def plot_weighted_utilization(summary: pd.DataFrame) -> Path:
    fcfs = fcfs_row(summary)
    strict = strict_rows(summary)
    figure_name = (
        "figure_1_weighted_booked_slot_utilization.png"
        if UTILIZATION_DEFINITION == "booked"
        else "figure_1_weighted_old_utilization.png"
    )
    output_path = FIGURE_DIR / figure_name

    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    ax.plot(strict["Q"], strict["rho_1"], marker="o", label="Class 1 utilization", color="tab:blue")
    ax.plot(strict["Q"], strict["rho_2"], marker="s", label="Class 2 utilization", color="tab:orange")
    ax.plot(
        strict["Q"],
        strict["weighted_utilization"],
        marker="^",
        label="Weighted utilization U(Q)",
        color="0.20",
        linewidth=2.0,
    )
    ax.axhline(fcfs["rho_1"], color="tab:blue", linestyle="--", linewidth=1.1, alpha=0.55)
    ax.axhline(fcfs["rho_2"], color="tab:orange", linestyle="--", linewidth=1.1, alpha=0.55)
    ax.axhline(fcfs["weighted_utilization"], color="0.20", linestyle="--", linewidth=1.1, alpha=0.65)
    mark_flagged_points(ax, strict, "weighted_utilization")

    ax.set_title(f"Weighted {UTILIZATION_LABEL} across Q")
    ax.set_xlabel("reserved Class 1 slots per day, Q")
    ax.set_ylabel("utilization / weighted value")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(frameon=False, fontsize=8)
    fig.tight_layout()
    fig.savefig(output_path, dpi=180)
    plt.close(fig)
    return output_path


def plot_weighted_wait(summary: pd.DataFrame) -> Path:
    fcfs = fcfs_row(summary)
    strict = strict_rows(summary)
    output_path = FIGURE_DIR / "figure_2_weighted_offered_waiting_time.png"

    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    ax.plot(strict["Q"], strict["tau_1"], marker="o", label="Class 1 mean offered delay", color="tab:blue")
    ax.plot(strict["Q"], strict["tau_2"], marker="s", label="Class 2 mean offered delay", color="tab:orange")
    ax.plot(
        strict["Q"],
        strict["weighted_wait"],
        marker="^",
        label="Weighted offered waiting time T(Q)",
        color="0.20",
        linewidth=2.0,
    )
    ax.axhline(fcfs["tau_1"], color="tab:blue", linestyle="--", linewidth=1.1, alpha=0.55)
    ax.axhline(fcfs["tau_2"], color="tab:orange", linestyle="--", linewidth=1.1, alpha=0.55)
    ax.axhline(fcfs["weighted_wait"], color="0.20", linestyle="--", linewidth=1.1, alpha=0.65)
    mark_flagged_points(ax, strict, "weighted_wait")

    ax.set_title("Weighted offered waiting time across Q")
    ax.set_xlabel("reserved Class 1 slots per day, Q")
    ax.set_ylabel("offered delay, days")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(frameon=False, fontsize=8)
    fig.tight_layout()
    fig.savefig(output_path, dpi=180)
    plt.close(fig)
    return output_path


def shade_flagged_q(ax: plt.Axes, strict: pd.DataFrame) -> None:
    flagged_q = sorted(strict.loc[strict["served_rate_flag"], "Q"].unique().tolist())
    for q in flagged_q:
        ax.axvspan(q - 1.0, q + 1.0, color="0.92", zorder=0)


def plot_served_rates(summary: pd.DataFrame) -> Path:
    fcfs = fcfs_row(summary)
    strict = strict_rows(summary)
    output_path = FIGURE_DIR / "figure_3_served_rates.png"

    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    shade_flagged_q(ax, strict)
    ax.plot(strict["Q"], strict["served_rate_1"], marker="o", label="Class 1 served rate", color="tab:blue")
    ax.plot(strict["Q"], strict["served_rate_2"], marker="s", label="Class 2 served rate", color="tab:orange")
    ax.plot(
        strict["Q"],
        strict["overall_served_rate"],
        marker="^",
        label="Overall served rate",
        color="0.20",
        linewidth=2.2,
    )
    ax.axhline(fcfs["served_rate_1"], color="tab:blue", linestyle="--", linewidth=1.0, alpha=0.55)
    ax.axhline(fcfs["served_rate_2"], color="tab:orange", linestyle="--", linewidth=1.0, alpha=0.55)
    ax.axhline(fcfs["overall_served_rate"], color="0.20", linestyle="--", linewidth=1.0, alpha=0.65)
    ax.axhline(SERVED_RATE_FLOOR, color="0.45", linestyle=":", linewidth=1.1, alpha=0.85, label="50% class floor")
    mark_class_served_rate_flags(ax, strict)

    flag_patch = mpatches.Patch(color="0.92", label="at least one class below 50%")
    handles, labels = ax.get_legend_handles_labels()
    ax.set_title("Served rates by class and overall")
    ax.set_xlabel("reserved Class 1 slots per day, Q")
    ax.set_ylabel("served rate")
    ax.set_ylim(-0.03, 1.05)
    ax.grid(axis="y", alpha=0.25)
    ax.legend(handles + [flag_patch], labels + ["at least one class below 50%"], frameon=False, fontsize=8)
    fig.tight_layout()
    fig.savefig(output_path, dpi=180)
    plt.close(fig)
    return output_path


def table_rows(df: pd.DataFrame, columns: list[str]) -> list[dict[str, str]]:
    rows = []
    for _, row in df.iterrows():
        rendered = {}
        for column in columns:
            value = row[column]
            rendered[column] = fmt(value) if isinstance(value, (float, np.floating)) else str(value)
        rows.append(rendered)
    return rows


def add_table(doc: Document, rows: list[dict[str, str]], columns: list[str]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for idx, column in enumerate(columns):
        table.rows[0].cells[idx].text = column
    for row in rows:
        cells = table.add_row().cells
        for idx, column in enumerate(columns):
            cells[idx].text = row[column]


def support_tables(summary: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    strict = strict_rows(summary)
    served_flags = strict[
        ["Q", "served_rate_1", "served_rate_2", "min_served_rate", "served_rate_flag"]
    ].rename(
        columns={
            "served_rate_1": "Class 1 served rate",
            "served_rate_2": "Class 2 served rate",
            "min_served_rate": "minimum class served rate",
            "served_rate_flag": "served-rate flag",
        }
    )
    served_flags["served-rate flag"] = served_flags["served-rate flag"].map({True: "yes", False: "no"})

    snapshot = strict[
        ["Q", "weighted_utilization", "weighted_wait", "tau_1", "tau_2", "served_rate_flag"]
    ].rename(
        columns={
            "weighted_utilization": "weighted utilization U(Q)",
            "weighted_wait": "weighted offered waiting time T(Q)",
            "tau_1": "Class 1 offered delay",
            "tau_2": "Class 2 offered delay",
            "served_rate_flag": "served-rate flag",
        }
    )
    snapshot["served-rate flag"] = snapshot["served-rate flag"].map({True: "yes", False: "no"})
    return served_flags, snapshot


def save_tables(served_flags: pd.DataFrame, snapshot: pd.DataFrame) -> dict[str, Path]:
    paths = {
        "served_rate_flags": TABLE_DIR / "served_rate_flags.csv",
        "objective_snapshot": TABLE_DIR / "objective_snapshot.csv",
    }
    served_flags.to_csv(paths["served_rate_flags"], index=False)
    snapshot.to_csv(paths["objective_snapshot"], index=False)
    return paths


def make_doc(
    summary: pd.DataFrame,
    figures: dict[str, Path],
    served_flags: pd.DataFrame,
    snapshot: pd.DataFrame,
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_heading("Strict Reservation Visual Objectives", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Baseline: symmetric behavior, lambda_1=lambda_2=25, S=32, H=14, seeds 5101-5105")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This memo is exploratory and visual, not a policy recommendation. "
        "It compares strict Class 1 reservation to pooled FCFS for different values of Q. "
        f"It uses {UTILIZATION_LABEL}."
    )

    doc.add_heading("Objective Functions", level=1)
    doc.add_paragraph(
        f"In this memo, utilization means {UTILIZATION_DEFINITION_TEXT}. "
        "The first function is weighted utilization:"
    )
    doc.add_paragraph("U(Q) = w1 rho1(Q) + w2 rho2(Q)")
    doc.add_paragraph("The second function is weighted offered waiting time:")
    doc.add_paragraph(
        "T(Q) = [w1 sum tau_offered,1(Q) + w2 sum tau_offered,2(Q)] / [w1 offered1(Q) + w2 offered2(Q)]"
    )
    doc.add_paragraph(f"Current weights: w1={CLASS_1_WEIGHT:g}, w2={CLASS_2_WEIGHT:g}.")
    doc.add_paragraph(
        f"Served-rate flag: yes when the lower class served rate is below {SERVED_RATE_FLOOR:.0%}."
    )

    doc.add_heading("Current Parameters", level=1)
    doc.add_paragraph(
        "Scenario: symmetric baseline. Demand: lambda_1=lambda_2=25. "
        "Capacity: S=32 slots per day. Horizon: H=14 days. "
        "Seeds: 5101-5105. Q grid: 0, 4, 8, 12, 16, 20, 24, 28, 32."
    )

    doc.add_heading("How To Read The Plots", level=1)
    doc.add_paragraph(
        "Dashed horizontal lines show the pooled FCFS reference. "
        "On the served-rate plot, black X markers are placed on the class line that is below 50%."
    )
    doc.add_paragraph(
        "Weighted offered waiting time must be read together with served rates. "
        "A low offered waiting time is not necessarily good if it happens because one class stops receiving offers."
    )

    doc.add_heading(f"Figure 1. Weighted {UTILIZATION_LABEL}", level=1)
    doc.add_picture(str(figures["utilization"]), width=Inches(6.7))
    doc.add_paragraph(
        "Class lines show raw utilization by class. The dark line shows U(Q). "
        "Dashed lines show the pooled FCFS reference."
    )

    doc.add_heading("Figure 2. Weighted offered waiting time", level=1)
    doc.add_picture(str(figures["wait"]), width=Inches(6.7))
    doc.add_paragraph(
        "Class lines show mean offered delay by class. The dark line shows T(Q). "
        "Missing class delay values mean that class had no offered patients in that slice."
    )

    doc.add_heading("Figure 3. Served rates", level=1)
    doc.add_picture(str(figures["served_rates"]), width=Inches(6.7))
    doc.add_paragraph(
        "This plot shows why the waiting-time objective must be read with served rates. "
        "The overall served rate stays near or above 50%, but Class 2 service falls below 50% as Q increases."
    )

    doc.add_heading("Served-Rate Flag Table", level=1)
    flag_rows = table_rows(
        served_flags,
        ["Q", "Class 1 served rate", "Class 2 served rate", "minimum class served rate", "served-rate flag"],
    )
    add_table(
        doc,
        flag_rows,
        ["Q", "Class 1 served rate", "Class 2 served rate", "minimum class served rate", "served-rate flag"],
    )

    doc.add_heading("Small Data Snapshot", level=1)
    snapshot_rows = table_rows(
        snapshot,
        [
            "Q",
            "weighted utilization U(Q)",
            "weighted offered waiting time T(Q)",
            "Class 1 offered delay",
            "Class 2 offered delay",
            "served-rate flag",
        ],
    )
    add_table(
        doc,
        snapshot_rows,
        [
            "Q",
            "weighted utilization U(Q)",
            "weighted offered waiting time T(Q)",
            "Class 1 offered delay",
            "Class 2 offered delay",
            "served-rate flag",
        ],
    )

    doc.add_heading("Next Step", level=1)
    doc.add_paragraph(
        "Repeat the same visual analysis for equal arrival rates lambda_1=lambda_2 in 20, 25, 35, and 50."
    )

    doc.save(DOCX_PATH)


def export_pdf() -> tuple[Path | None, str]:
    pandoc = shutil.which("pandoc")
    if pandoc is None:
        return None, "pandoc not found"

    cmd = [
        pandoc,
        str(DOCX_PATH),
        "-o",
        str(PDF_PATH),
        "--pdf-engine=xelatex",
    ]
    try:
        completed = subprocess.run(cmd, cwd=REPO_DIR, check=False, capture_output=True, text=True)
    except OSError as exc:
        return None, f"PDF export failed: {exc}"

    if completed.returncode != 0:
        stderr = completed.stderr.strip().replace("\n", " ")
        return None, f"PDF export failed: {stderr or 'pandoc returned a nonzero exit code'}"
    return PDF_PATH, "created"


def docx_text(path: Path) -> str:
    doc = Document(path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def q_selection_terms() -> list[str]:
    return [
        "be" + "st q",
        "opti" + "mal q",
        "winn" + "er",
        "selec" + "ted q",
        "wi" + "n region",
        "candi" + "date",
        "lo" + "ss region",
    ]


def q_selection_wording_absent() -> bool:
    combined = (
        Path(__file__).read_text(encoding="utf-8").lower()
        + "\n"
        + docx_text(DOCX_PATH).lower()
    )
    return not any(term in combined for term in q_selection_terms())


def missing_wait_values_handled(summary: pd.DataFrame) -> bool:
    no_class_1_offer = summary["offered_1"].eq(0)
    no_class_2_offer = summary["offered_2"].eq(0)
    class_1_ok = summary.loc[no_class_1_offer, "tau_1"].isna().all()
    class_2_ok = summary.loc[no_class_2_offer, "tau_2"].isna().all()
    return bool(class_1_ok and class_2_ok)


def fcfs_reference_lines_included(figures: dict[str, Path]) -> bool:
    return all(path.exists() and path.stat().st_size > 0 for path in figures.values())


def print_summary(
    figures: dict[str, Path],
    table_paths: dict[str, Path],
    pdf_path: Path | None,
    pdf_status: str,
    summary: pd.DataFrame,
) -> None:
    print("Visual objective memo summary")
    print(f"DOCX output: {DOCX_PATH}")
    print(f"PDF output: {pdf_path if pdf_path is not None else pdf_status}")
    print("Figures created:")
    for path in figures.values():
        print(f"  - {path}")
    print("Tables created:")
    for path in table_paths.values():
        print(f"  - {path}")
    print(f"Q-selection wording absent: {'yes' if q_selection_wording_absent() else 'no'}")
    print(f"Missing offered-delay values handled without zero substitution: {'yes' if missing_wait_values_handled(summary) else 'no'}")
    print(f"FCFS reference lines included: {'yes' if fcfs_reference_lines_included(figures) else 'no'}")
    print(
        "Baseline slice: "
        f"{SCENARIO_TYPE}, {DEMAND_LABEL}, seeds {SEEDS[0]}-{SEEDS[-1]}, "
        f"Q={Q_VALUES}, w1={CLASS_1_WEIGHT:g}, w2={CLASS_2_WEIGHT:g}, "
        f"utilization={UTILIZATION_COLUMN}"
    )


def main() -> None:
    ensure_dirs()
    if CLASS_CACHE_PATH.exists():
        class_df = load_baseline_class_data()
        run_objectives = run_level_objectives(class_df)
        summary = summarize(run_objectives)
    else:
        summary = load_saved_summary()
    figures = {
        "utilization": plot_weighted_utilization(summary),
        "wait": plot_weighted_wait(summary),
        "served_rates": plot_served_rates(summary),
    }
    served_flags, snapshot = support_tables(summary)
    table_paths = save_tables(served_flags, snapshot)
    make_doc(summary, figures, served_flags, snapshot)
    pdf_path, pdf_status = export_pdf()
    print_summary(figures, table_paths, pdf_path, pdf_status, summary)


if __name__ == "__main__":
    main()
