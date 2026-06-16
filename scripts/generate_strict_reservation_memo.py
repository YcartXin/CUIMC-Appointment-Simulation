from __future__ import annotations

import json
import math
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import matplotlib.pyplot as plt
import matplotlib.patheffects as path_effects
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


REPO_DIR = Path(__file__).resolve().parents[1]
OUTPUT_DIR = REPO_DIR / "outputs" / "reservation_report"
FIGURE_DIR = OUTPUT_DIR / "figures"
TABLE_DIR = OUTPUT_DIR / "tables"
DOCX_PATH = OUTPUT_DIR / "strict_reservation_util_wait_tradeoff_memo.docx"
PDF_PATH = OUTPUT_DIR / "strict_reservation_util_wait_tradeoff_memo.pdf"

CACHE_DIR = REPO_DIR / "notebooks" / "temp" / "reservation_playground_cache" / "92a74cbf5416"
RUN_CACHE_PATH = CACHE_DIR / "run_df.pkl"
PARAMS_CACHE_PATH = CACHE_DIR / "simulation_params.json"
EXISTING_DOCX_PATH = REPO_DIR / "notebooks" / "reservation_objective_analysis_summary.docx"

STRICT_POLICY = "Strict C1 reservation"
FCFS_POLICY = "Pooled FCFS"
SCENARIO_ORDER = ["symmetric_baseline", "class_1_advantaged", "class_1_disadvantaged"]
SCENARIO_LABELS = {
    "symmetric_baseline": "Symmetric baseline",
    "class_1_advantaged": "Class 1 advantaged",
    "class_1_disadvantaged": "Class 1 disadvantaged",
}

REQUESTED_Q_VALUES = [0, 4, 8, 12, 16, 20, 24, 28, 32]
HEADLINE_W1 = 1.5
HEADLINE_W2 = 1.0
HEADLINE_C = 0.05
HEADLINE_GAMMA = 0.05
HEADLINE_SERVED_RATE_MIN = 0.55
EPS = 1e-9

W1_VALUES = [1.0, 1.25, 1.5, 2.0, 3.0]
C_VALUES = [0.0, 0.02, 0.05, 0.10]
GAMMA_VALUES = [0.0, 0.02, 0.05, 0.10]
SERVED_RATE_MIN_VALUES = [0.45, 0.55, 0.65, 0.70]


@dataclass(frozen=True)
class AuditResult:
    old_doc_terms_found: bool
    old_plot_terms_found: bool
    booked_metric_available: bool
    offered_delay_available: bool
    service_rates_available: bool


def safe_divide(numerator: float, denominator: float) -> float:
    return numerator / denominator if denominator else 0.0


def ensure_dirs() -> None:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    TABLE_DIR.mkdir(parents=True, exist_ok=True)


def audit_inputs(run_df: pd.DataFrame) -> AuditResult:
    old_doc_terms_found = False
    if EXISTING_DOCX_PATH.exists():
        try:
            from docx import Document as ReadDocument

            old_doc = ReadDocument(EXISTING_DOCX_PATH)
            text = "\n".join(paragraph.text for paragraph in old_doc.paragraphs).lower()
            old_doc_terms_found = (
                "weighted served" in text
                or "net priority" in text
                or "net_priority_score" in text
            )
        except Exception:
            old_doc_terms_found = False

    old_plot_terms_found = any(
        "net_priority" in path.name
        for path in (REPO_DIR / "notebooks" / "temp" / "objective_doc_figures_selected").glob("*.png")
    )
    booked_metric_available = "booked_slot_utilization" in run_df.columns
    offered_delay_available = all(
        column in run_df.columns
        for column in [
            "class_1_total_offered_delay",
            "class_2_total_offered_delay",
            "class_1_offered",
            "class_2_offered",
        ]
    )
    service_rates_available = all(
        column in run_df.columns
        for column in ["class_1_served_rate", "class_2_served_rate", "min_class_served_rate"]
    )
    return AuditResult(
        old_doc_terms_found=old_doc_terms_found,
        old_plot_terms_found=old_plot_terms_found,
        booked_metric_available=booked_metric_available,
        offered_delay_available=offered_delay_available,
        service_rates_available=service_rates_available,
    )


def load_run_data() -> tuple[pd.DataFrame, dict]:
    if not RUN_CACHE_PATH.exists():
        raise FileNotFoundError(f"Missing cached run data: {RUN_CACHE_PATH}")

    run_df = pd.read_pickle(RUN_CACHE_PATH)
    params = json.loads(PARAMS_CACHE_PATH.read_text()) if PARAMS_CACHE_PATH.exists() else {}
    run_df = run_df[
        (run_df["policy"] == FCFS_POLICY)
        | ((run_df["policy"] == STRICT_POLICY) & run_df["Q"].isin(REQUESTED_Q_VALUES))
    ].copy()
    run_df["scenario_label"] = run_df["scenario_type"].map(SCENARIO_LABELS)
    run_df["total_demand_label"] = run_df["lambda_total"].map(lambda value: f"lambda={value:g}")
    return run_df, params


def weighted_delay(row: pd.Series, w1: float, w2: float = HEADLINE_W2) -> float:
    return safe_divide(
        w1 * row["class_1_total_offered_delay"] + w2 * row["class_2_total_offered_delay"],
        w1 * row["class_1_offered"] + w2 * row["class_2_offered"],
    )


def with_objectives(
    run_df: pd.DataFrame,
    *,
    w1: float,
    c: float,
    gamma: float,
) -> pd.DataFrame:
    df = run_df.copy()
    df["study_a_objective"] = df["booked_slot_utilization"]
    df["weighted_booked"] = (
        w1 * df["class_1_booked_slot_utilization"]
        + HEADLINE_W2 * df["class_2_booked_slot_utilization"]
    )
    df["tau_bar_w"] = df.apply(lambda row: weighted_delay(row, w1, HEADLINE_W2), axis=1)
    df["study_b_objective"] = (
        df["weighted_booked"]
        - c * df["Q"].div(df["slots_per_day"])
        - gamma * df["tau_bar_w"].div(df["horizon_days"])
    )
    return df


def paired_summary(
    run_df: pd.DataFrame,
    *,
    objective_column: str,
    served_rate_min: float,
) -> pd.DataFrame:
    pair_keys = ["scenario_id", "scenario_type", "scenario_label", "lambda_total", "seed"]
    metric_cols = [
        objective_column,
        "booked_slot_utilization",
        "average_utilization",
        "class_1_served_rate",
        "class_2_served_rate",
        "min_class_served_rate",
        "mean_offered_booking_delay",
    ]

    fcfs = (
        run_df[run_df["policy"] == FCFS_POLICY][pair_keys + metric_cols]
        .rename(columns={column: f"fcfs_{column}" for column in metric_cols})
    )
    strict = run_df[run_df["policy"] == STRICT_POLICY].copy()
    paired = strict.merge(fcfs, on=pair_keys, how="left")
    if paired[f"fcfs_{objective_column}"].isna().any():
        raise RuntimeError(f"Missing FCFS matches for {objective_column}")

    paired["delta_objective"] = paired[objective_column] - paired[f"fcfs_{objective_column}"]
    for column in metric_cols:
        paired[f"delta_{column}"] = paired[column] - paired[f"fcfs_{column}"]
    paired["feasible_seed"] = paired["min_class_served_rate"] >= served_rate_min

    group_keys = ["scenario_id", "scenario_type", "scenario_label", "lambda_total", "Q"]
    summary = (
        paired.groupby(group_keys)
        .agg(
            mean_delta=("delta_objective", "mean"),
            std_delta=("delta_objective", "std"),
            count=("delta_objective", "count"),
            strict_objective=(objective_column, "mean"),
            fcfs_objective=(f"fcfs_{objective_column}", "mean"),
            booked_utilization=("booked_slot_utilization", "mean"),
            fcfs_booked_utilization=("fcfs_booked_slot_utilization", "mean"),
            delta_booked_utilization=("delta_booked_slot_utilization", "mean"),
            attended_utilization=("average_utilization", "mean"),
            fcfs_attended_utilization=("fcfs_average_utilization", "mean"),
            delta_attended_utilization=("delta_average_utilization", "mean"),
            class_1_served_rate=("class_1_served_rate", "mean"),
            class_2_served_rate=("class_2_served_rate", "mean"),
            min_class_served_rate=("min_class_served_rate", "mean"),
            delta_class_1_served_rate=("delta_class_1_served_rate", "mean"),
            delta_class_2_served_rate=("delta_class_2_served_rate", "mean"),
            mean_offered_booking_delay=("mean_offered_booking_delay", "mean"),
            fcfs_mean_offered_booking_delay=("fcfs_mean_offered_booking_delay", "mean"),
            delta_mean_offered_booking_delay=("delta_mean_offered_booking_delay", "mean"),
            feasible_seed_share=("feasible_seed", "mean"),
        )
        .reset_index()
    )
    summary["sem_delta"] = summary["std_delta"].fillna(0.0).div(np.sqrt(summary["count"]))
    summary["ci_low"] = summary["mean_delta"] - 1.96 * summary["sem_delta"]
    summary["ci_high"] = summary["mean_delta"] + 1.96 * summary["sem_delta"]
    summary["feasible"] = summary["min_class_served_rate"] >= served_rate_min
    summary["classification"] = np.select(
        [
            (summary["mean_delta"] > EPS) & (~summary["feasible"]),
            (summary["mean_delta"] > EPS) & (summary["feasible"]) & (summary["ci_low"] <= EPS),
            (summary["mean_delta"] > EPS) & (summary["feasible"]) & (summary["ci_low"] > EPS),
        ],
        ["infeasible utilization win", "possible win", "feasible win"],
        default="loss",
    )
    return summary.sort_values(["scenario_type", "lambda_total", "Q"])


def choose_best_q(summary: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for _, group in summary.groupby(["scenario_id", "scenario_type", "scenario_label", "lambda_total"]):
        feasible = group[group["feasible"]]
        candidates = feasible if not feasible.empty else group
        sortable = candidates.copy()
        sortable["sort_objective"] = sortable["strict_objective"].round(12)
        sortable["sort_delta"] = sortable["mean_delta"].where(
            sortable["mean_delta"].abs() > EPS,
            0.0,
        )
        best = sortable.sort_values(
            ["sort_objective", "sort_delta", "Q"],
            ascending=[False, False, True],
        ).iloc[0]
        rows.append(best)
    return pd.DataFrame(rows).reset_index(drop=True)


def binding_reason(row: pd.Series) -> str:
    if not bool(row["feasible"]):
        return "service-rate guardrail"
    if row["mean_delta"] <= EPS:
        return "objective does not beat FCFS"
    return "none"


def baseline_policy_table(
    scored: pd.DataFrame,
    best_a: pd.DataFrame,
    best_b: pd.DataFrame,
    *,
    served_rate_min: float,
) -> pd.DataFrame:
    baseline_filter = (
        (scored["scenario_type"] == "symmetric_baseline")
        & (scored["lambda_total"] == 50.0)
    )
    best_a_q = int(best_a.loc[
        (best_a["scenario_type"] == "symmetric_baseline") & (best_a["lambda_total"] == 50.0),
        "Q",
    ].iloc[0])
    best_b_q = int(best_b.loc[
        (best_b["scenario_type"] == "symmetric_baseline") & (best_b["lambda_total"] == 50.0),
        "Q",
    ].iloc[0])

    rows = []
    selections = [(FCFS_POLICY, 0, "Pooled FCFS")]
    if best_a_q == best_b_q:
        selections.append(
            (
                STRICT_POLICY,
                best_a_q,
                f"Study A/B selected Q={best_a_q} (same booking rule as FCFS)" if best_a_q == 0 else f"Strict reservation, Study A/B best Q={best_a_q}",
            )
        )
    else:
        selections.extend(
            [
                (STRICT_POLICY, best_a_q, f"Strict reservation, Study A best Q={best_a_q}"),
                (STRICT_POLICY, best_b_q, f"Strict reservation, Study B best Q={best_b_q}"),
            ]
        )
    if best_a_q == 0 and best_b_q == 0:
        selections.append((STRICT_POLICY, 4, "Strict reservation, smallest positive Q=4"))

    seen = set()
    for policy, q, label in selections:
        key = (policy, q)
        if key in seen:
            continue
        seen.add(key)
        group = scored[baseline_filter & (scored["policy"] == policy) & (scored["Q"] == q)]
        rows.append(
            {
                "policy": label,
                "Q": q,
                "booked utilization": group["booked_slot_utilization"].mean(),
                "attended utilization": group["average_utilization"].mean(),
                "Class 1 served rate": group["class_1_served_rate"].mean(),
                "Class 2 served rate": group["class_2_served_rate"].mean(),
                "minimum class served rate": group["min_class_served_rate"].mean(),
                "average offered delay": group["mean_offered_booking_delay"].mean(),
                "Study A objective": group["study_a_objective"].mean(),
                "Study B objective": group["study_b_objective"].mean(),
                "feasible under served-rate threshold?": (
                    "yes" if group["min_class_served_rate"].mean() >= served_rate_min else "no"
                ),
            }
        )
    return pd.DataFrame(rows)


def best_q_comparison_table(best_a: pd.DataFrame, best_b: pd.DataFrame) -> pd.DataFrame:
    left = best_a[
        [
            "scenario_type",
            "scenario_label",
            "lambda_total",
            "Q",
            "feasible",
            "mean_delta",
            "delta_class_1_served_rate",
            "delta_class_2_served_rate",
            "delta_booked_utilization",
            "delta_mean_offered_booking_delay",
        ]
    ].rename(
        columns={
            "Q": "best Q under Study A",
            "feasible": "Study A feasible?",
            "mean_delta": "Study A strict-minus-FCFS difference",
            "delta_class_1_served_rate": "Study A Class 1 served-rate change",
            "delta_class_2_served_rate": "Study A Class 2 served-rate change",
            "delta_booked_utilization": "Study A utilization change",
            "delta_mean_offered_booking_delay": "Study A wait change",
        }
    )
    right = best_b[
        [
            "scenario_type",
            "lambda_total",
            "Q",
            "feasible",
            "mean_delta",
            "delta_class_1_served_rate",
            "delta_class_2_served_rate",
            "delta_booked_utilization",
            "delta_mean_offered_booking_delay",
        ]
    ].rename(
        columns={
            "Q": "best Q under Study B",
            "feasible": "Study B feasible?",
            "mean_delta": "Study B strict-minus-FCFS difference",
            "delta_class_1_served_rate": "Study B Class 1 served-rate change",
            "delta_class_2_served_rate": "Study B Class 2 served-rate change",
            "delta_booked_utilization": "Study B utilization change",
            "delta_mean_offered_booking_delay": "Study B wait change",
        }
    )
    table = left.merge(right, on=["scenario_type", "lambda_total"], how="inner")
    table["scenario"] = table.apply(
        lambda row: f"{row['scenario_label']}, lambda={row['lambda_total']:g}",
        axis=1,
    )
    table["do studies agree?"] = np.where(
        table["best Q under Study A"] == table["best Q under Study B"],
        "yes",
        "no",
    )
    table["main reason if they differ"] = np.where(
        table["do studies agree?"] == "yes",
        "same selected Q",
        "wait/cost penalty changes ranking",
    )
    return table[
        [
            "scenario",
            "best Q under Study A",
            "Study A feasible?",
            "best Q under Study B",
            "Study B feasible?",
            "do studies agree?",
            "main reason if they differ",
            "Study A strict-minus-FCFS difference",
            "Study B strict-minus-FCFS difference",
            "Study A Class 1 served-rate change",
            "Study A Class 2 served-rate change",
            "Study B Class 1 served-rate change",
            "Study B Class 2 served-rate change",
            "Study A utilization change",
            "Study B utilization change",
            "Study A wait change",
            "Study B wait change",
        ]
    ]


def baseline_sensitivity(scored: pd.DataFrame) -> pd.DataFrame:
    rows = []
    baseline = scored[
        (scored["scenario_type"] == "symmetric_baseline")
        & (scored["lambda_total"] == 50.0)
    ].copy()

    for served_min in SERVED_RATE_MIN_VALUES:
        a_summary = paired_summary(baseline, objective_column="study_a_objective", served_rate_min=served_min)
        best_a = choose_best_q(a_summary).iloc[0]
        rows.append(
            {
                "study": "A: utilization-first",
                "w1": 1.0,
                "c": 0.0,
                "gamma": 0.0,
                "served_rate_min": served_min,
                "best Q": int(best_a["Q"]),
                "strict beats FCFS?": bool(best_a["mean_delta"] > EPS and best_a["feasible"]),
                "binding constraint or main failure mode": binding_reason(best_a),
                "mean strict-minus-FCFS objective": best_a["mean_delta"],
            }
        )

    for w1 in W1_VALUES:
        for c in C_VALUES:
            for gamma in GAMMA_VALUES:
                varied = with_objectives(baseline, w1=w1, c=c, gamma=gamma)
                for served_min in SERVED_RATE_MIN_VALUES:
                    b_summary = paired_summary(
                        varied,
                        objective_column="study_b_objective",
                        served_rate_min=served_min,
                    )
                    best_b = choose_best_q(b_summary).iloc[0]
                    rows.append(
                        {
                            "study": "B: wait-adjusted",
                            "w1": w1,
                            "c": c,
                            "gamma": gamma,
                            "served_rate_min": served_min,
                            "best Q": int(best_b["Q"]),
                            "strict beats FCFS?": bool(best_b["mean_delta"] > EPS and best_b["feasible"]),
                            "binding constraint or main failure mode": binding_reason(best_b),
                            "mean strict-minus-FCFS objective": best_b["mean_delta"],
                        }
                    )

    return pd.DataFrame(rows)


def save_tables(
    baseline_table: pd.DataFrame,
    best_table: pd.DataFrame,
    sensitivity_table: pd.DataFrame,
    study_a_summary: pd.DataFrame,
    study_b_summary: pd.DataFrame,
) -> None:
    baseline_table.to_csv(TABLE_DIR / "baseline_policy_comparison.csv", index=False)
    best_table.to_csv(TABLE_DIR / "best_q_by_study.csv", index=False)
    sensitivity_table.to_csv(TABLE_DIR / "sensitivity_summary.csv", index=False)
    study_a_summary.to_csv(TABLE_DIR / "study_a_utilization_first_summary.csv", index=False)
    study_b_summary.to_csv(TABLE_DIR / "study_b_wait_adjusted_summary.csv", index=False)


def heatmap_plot(
    summary: pd.DataFrame,
    *,
    value: str,
    title: str,
    output_path: Path,
    served_rate_min: float,
) -> None:
    fig, axes = plt.subplots(1, 3, figsize=(13.2, 4.1), sharey=True)
    max_abs = float(np.nanmax(np.abs(summary[value].values))) or 1.0
    for ax, scenario_type in zip(axes, SCENARIO_ORDER):
        subset = summary[summary["scenario_type"] == scenario_type]
        pivot = subset.pivot(index="lambda_total", columns="Q", values=value).sort_index()
        feasible = subset.pivot(index="lambda_total", columns="Q", values="feasible").reindex_like(pivot)
        image = ax.imshow(pivot.values, aspect="auto", cmap="RdBu", vmin=-max_abs, vmax=max_abs)
        ax.set_title(SCENARIO_LABELS[scenario_type])
        ax.set_xticks(range(len(pivot.columns)))
        ax.set_xticklabels([int(q) for q in pivot.columns], rotation=45)
        ax.set_yticks(range(len(pivot.index)))
        ax.set_yticklabels([int(value) for value in pivot.index])
        ax.set_xlabel("reserved slots per day, Q")
        for y_idx in range(feasible.shape[0]):
            for x_idx in range(feasible.shape[1]):
                if not bool(feasible.iloc[y_idx, x_idx]):
                    ax.text(x_idx, y_idx, "x", ha="center", va="center", color="black", fontsize=10, fontweight="bold")
    axes[0].set_ylabel("total demand, lambda")
    cbar = fig.colorbar(image, ax=axes.ravel().tolist(), shrink=0.88)
    cbar.set_label("strict reservation minus FCFS")
    fig.suptitle(f"{title}\nblack x: strict min class served rate < {served_rate_min:.2f}", y=1.02)
    fig.savefig(output_path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def service_rate_plot(scored: pd.DataFrame, output_path: Path, served_rate_min: float) -> None:
    baseline = scored[
        (scored["scenario_type"] == "symmetric_baseline")
        & (scored["lambda_total"] == 50.0)
    ]
    strict = (
        baseline[baseline["policy"] == STRICT_POLICY]
        .groupby("Q")[["class_1_served_rate", "class_2_served_rate"]]
        .mean()
        .reset_index()
        .sort_values("Q")
    )
    fcfs = baseline[baseline["policy"] == FCFS_POLICY]
    fcfs_c1 = fcfs["class_1_served_rate"].mean()
    fcfs_c2 = fcfs["class_2_served_rate"].mean()

    fig, ax = plt.subplots(figsize=(7.7, 4.2))
    ax.plot(strict["Q"], strict["class_1_served_rate"], marker="o", label="Strict: Class 1")
    ax.plot(strict["Q"], strict["class_2_served_rate"], marker="o", label="Strict: Class 2")
    ax.axhline(fcfs_c1, color="tab:blue", linestyle="--", linewidth=1.1, alpha=0.65, label="FCFS Class 1")
    ax.axhline(fcfs_c2, color="tab:orange", linestyle="--", linewidth=1.1, alpha=0.65, label="FCFS Class 2")
    ax.axhline(served_rate_min, color="black", linestyle=":", linewidth=1.4, label=f"threshold {served_rate_min:.2f}")
    ax.set_title("Service-rate guardrail at baseline: symmetric, lambda=50")
    ax.set_xlabel("reserved slots per day, Q")
    ax.set_ylabel("served rate")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(frameon=False, ncol=2)
    fig.tight_layout()
    fig.savefig(output_path, dpi=180)
    plt.close(fig)


def best_q_plot(best_table: pd.DataFrame, output_path: Path) -> None:
    plot_df = best_table.copy()
    fig, axes = plt.subplots(1, 3, figsize=(13.0, 4.0), sharey=True)
    for ax, scenario_type in zip(axes, SCENARIO_ORDER):
        subset = plot_df[plot_df["scenario"].str.startswith(SCENARIO_LABELS[scenario_type])]
        x = [float(label.split("lambda=")[1]) for label in subset["scenario"]]
        ax.plot(x, subset["best Q under Study A"], marker="o", label="Study A")
        ax.plot(x, subset["best Q under Study B"], marker="s", label="Study B")
        for _, row in subset.iterrows():
            demand = float(row["scenario"].split("lambda=")[1])
            if not bool(row["Study A feasible?"]):
                marker = ax.scatter(
                    demand,
                    row["best Q under Study A"],
                    marker="X",
                    color="black",
                    edgecolors="white",
                    linewidths=1.2,
                    s=120,
                    zorder=10,
                )
                marker.set_path_effects([path_effects.withStroke(linewidth=2.5, foreground="white")])
            if not bool(row["Study B feasible?"]):
                marker = ax.scatter(
                    demand,
                    row["best Q under Study B"],
                    marker="X",
                    color="black",
                    edgecolors="white",
                    linewidths=1.2,
                    s=120,
                    zorder=10,
                )
                marker.set_path_effects([path_effects.withStroke(linewidth=2.5, foreground="white")])
        ax.set_title(SCENARIO_LABELS[scenario_type])
        ax.set_xlabel("total demand, lambda")
        ax.grid(axis="y", alpha=0.25)
    axes[0].set_ylabel("best reserved slots per day, Q")
    axes[0].legend(frameon=False)
    fig.suptitle("Best Q comparison: Study A utilization objective vs Study B wait-adjusted objective", y=1.02)
    fig.savefig(output_path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def selected_q_class_served_rate_plot(
    study_b_summary: pd.DataFrame,
    output_path: Path,
    served_rate_min: float,
) -> None:
    selected = choose_best_q(study_b_summary)
    fig, axes = plt.subplots(1, 3, figsize=(13.0, 4.0), sharey=True)
    for ax, scenario_type in zip(axes, SCENARIO_ORDER):
        subset = selected[selected["scenario_type"] == scenario_type].sort_values("lambda_total")
        ax.plot(
            subset["lambda_total"],
            subset["class_1_served_rate"],
            marker="o",
            label="Class 1",
            color="tab:blue",
        )
        ax.plot(
            subset["lambda_total"],
            subset["class_2_served_rate"],
            marker="s",
            label="Class 2",
            color="tab:orange",
        )
        ax.axhline(served_rate_min, color="black", linestyle=":", linewidth=1.3, label=f"threshold {served_rate_min:.2f}")
        for _, row in subset.iterrows():
            if not bool(row["feasible"]):
                for class_column in ["class_1_served_rate", "class_2_served_rate"]:
                    marker = ax.scatter(
                        row["lambda_total"],
                        row[class_column],
                        marker="X",
                        color="black",
                        edgecolors="white",
                        linewidths=1.2,
                        s=95,
                        zorder=10,
                    )
                    marker.set_path_effects([path_effects.withStroke(linewidth=2.5, foreground="white")])
        ax.set_title(SCENARIO_LABELS[scenario_type])
        ax.set_xlabel("total demand, lambda")
        ax.grid(axis="y", alpha=0.25)
    axes[0].set_ylabel("served rate at Study B selected Q")
    axes[0].legend(frameon=False)
    fig.suptitle("Class-specific served rates at the Study B selected Q", y=1.02)
    fig.savefig(output_path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def gamma_sensitivity_plot(sensitivity: pd.DataFrame, output_path: Path) -> None:
    subset = sensitivity[
        (sensitivity["study"] == "B: wait-adjusted")
        & (sensitivity["w1"] == HEADLINE_W1)
        & (sensitivity["served_rate_min"] == HEADLINE_SERVED_RATE_MIN)
    ].copy()
    fig, axes = plt.subplots(1, 2, figsize=(10.5, 4.0))
    for c, group in subset.groupby("c"):
        group = group.sort_values("gamma")
        axes[0].plot(group["gamma"], group["best Q"], marker="o", label=f"c={c:g}")
        axes[1].plot(group["gamma"], group["mean strict-minus-FCFS objective"], marker="o", label=f"c={c:g}")
    axes[0].set_title("Best Q")
    axes[0].set_ylabel("best Q")
    axes[1].set_title("Best strict-minus-FCFS objective")
    axes[1].set_ylabel("mean delta")
    for ax in axes:
        ax.set_xlabel("wait penalty gamma")
        ax.grid(axis="y", alpha=0.25)
        ax.axhline(0, color="0.25", linewidth=1)
    axes[1].legend(frameon=False, title="slot cost")
    fig.suptitle("Study B sensitivity at baseline: symmetric, lambda=50, w1=1.5", y=1.02)
    fig.tight_layout()
    fig.savefig(output_path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def generate_figures(
    scored: pd.DataFrame,
    study_a_summary: pd.DataFrame,
    study_b_summary: pd.DataFrame,
    best_table: pd.DataFrame,
    sensitivity: pd.DataFrame,
) -> dict[str, Path]:
    paths = {
        "study_a_heatmap": FIGURE_DIR / "study_a_utilization_heatmap.png",
        "study_b_heatmap": FIGURE_DIR / "study_b_wait_adjusted_heatmap.png",
        "service_rate": FIGURE_DIR / "service_rate_diagnostic_baseline.png",
        "best_q": FIGURE_DIR / "best_q_comparison.png",
        "selected_q_class_service": FIGURE_DIR / "selected_q_class_served_rates.png",
        "gamma_sensitivity": FIGURE_DIR / "gamma_sensitivity.png",
    }
    heatmap_plot(
        study_a_summary,
        value="mean_delta",
        title="Study A objective: strict-minus-FCFS booked-slot utilization",
        output_path=paths["study_a_heatmap"],
        served_rate_min=HEADLINE_SERVED_RATE_MIN,
    )
    heatmap_plot(
        study_b_summary,
        value="mean_delta",
        title="Study B objective: strict-minus-FCFS wait-adjusted booked-slot utility",
        output_path=paths["study_b_heatmap"],
        served_rate_min=HEADLINE_SERVED_RATE_MIN,
    )
    service_rate_plot(scored, paths["service_rate"], HEADLINE_SERVED_RATE_MIN)
    best_q_plot(best_table, paths["best_q"])
    selected_q_class_served_rate_plot(study_b_summary, paths["selected_q_class_service"], HEADLINE_SERVED_RATE_MIN)
    gamma_sensitivity_plot(sensitivity, paths["gamma_sensitivity"])
    return paths


def format_number(value: float, digits: int = 3) -> str:
    if pd.isna(value):
        return ""
    return f"{value:.{digits}f}"


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], *, digits: int = 3) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for idx, column in enumerate(columns):
        table.rows[0].cells[idx].text = column
    for _, row in df[columns].iterrows():
        cells = table.add_row().cells
        for idx, column in enumerate(columns):
            value = row[column]
            if isinstance(value, (float, np.floating)):
                cells[idx].text = format_number(float(value), digits)
            elif isinstance(value, (bool, np.bool_)):
                cells[idx].text = "yes" if value else "no"
            else:
                cells[idx].text = str(value)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def make_doc(
    *,
    baseline_table: pd.DataFrame,
    best_table: pd.DataFrame,
    sensitivity: pd.DataFrame,
    figure_paths: dict[str, Path],
    audit: AuditResult,
    seeds: list[int],
    params: dict,
) -> dict[str, object]:
    baseline_row_a = best_table[best_table["scenario"] == "Symmetric baseline, lambda=50"].iloc[0]
    study_a_best_q = int(baseline_row_a["best Q under Study A"])
    study_b_best_q = int(baseline_row_a["best Q under Study B"])
    wait_changed_best_q = study_a_best_q != study_b_best_q
    study_a_delta = float(baseline_row_a["Study A strict-minus-FCFS difference"])
    study_b_delta = float(baseline_row_a["Study B strict-minus-FCFS difference"])
    study_a_feasible = bool(baseline_row_a["Study A feasible?"])
    study_b_feasible = bool(baseline_row_a["Study B feasible?"])

    infeasible_util_wins = int(
        pd.read_csv(TABLE_DIR / "study_a_utilization_first_summary.csv")[
            lambda df: df["classification"].eq("infeasible utilization win")
        ].shape[0]
    )
    objectives_agree_count = int(best_table["do studies agree?"].eq("yes").sum())
    objective_rows = int(best_table.shape[0])

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = doc.add_heading("Strict Class 1 Reservation vs FCFS: Utilization, Service-Rate, and Waiting-Time Tradeoffs", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Short preliminary decision memo for review")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Executive Summary", level=1)
    add_bullets(
        doc,
        [
            "Bottom line: these 5-seed results are preliminary. Near the baseline, strict reservation does not beat FCFS on the utilization-first objective once the service-rate guardrail is checked.",
            f"Baseline setting: lambda_1=lambda_2=25, total lambda=50, S=32 slots/day, H=14 days, seeds {seeds[0]}-{seeds[-1]}.",
            f"Study A objective is U_util = rho^booked. At baseline it selects Q={study_a_best_q}; strict-minus-FCFS booked utilization is {study_a_delta:+.3f}; feasible={study_a_feasible}.",
            f"Study B objective is U_wait_util with w1=1.5, c=0.05, gamma=0.05. At baseline it selects Q={study_b_best_q}; strict-minus-FCFS utility is {study_b_delta:+.3f}; feasible={study_b_feasible}.",
            f"The wait penalty {'changes' if wait_changed_best_q else 'does not change'} the selected baseline Q relative to Study A.",
            "Service-rate guardrails matter: cells marked with x in the heatmaps fail min(served_rate_1, served_rate_2) >= 0.55 and should not be interpreted as operational wins.",
            f"Across the 12 demand/behavior scenarios, the two studies select the same Q in {objectives_agree_count} cases. Differences mostly occur when wait reduction and Class 1 priority favor large Q but the service-rate guardrail flags infeasibility.",
            "Decision point: decide whether booked-slot utilization is the right primary capacity target, or whether completed visits and Class 2 access should receive stronger constraints before choosing Q.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("1. What The Policy Changes", level=1)
    doc.add_paragraph(
        "Pooled FCFS uses one common appointment pool. Strict Class 1 reservation protects Q slots per day for Class 1. Class 1 can use protected slots first and then general slots. Class 2 can only use general slots. Unused protected slots remain empty."
    )
    doc.add_paragraph(
        "The objective functions in this memo are after-the-fact scoring rules. They do not change the simulation policy."
    )

    doc.add_heading("2. Why Utilization Alone Is Not Enough", level=1)
    doc.add_paragraph(
        "Booked utilization rho^booked measures whether service-day appointment slots were booked. It includes no-shows because those slots were occupied on the schedule. Attended utilization rho^attended measures completed visits. Served rate measures completed visits divided by arrivals."
    )
    doc.add_paragraph(
        "A high booked-utilization value can still be unacceptable if one class has poor service. For that reason, all best-Q decisions are checked against min(served_rate_1, served_rate_2) >= 0.55."
    )

    doc.add_heading("3. Study A: Utilization-First With Guardrails", level=1)
    doc.add_paragraph("Main objective used to choose Q:")
    doc.add_paragraph("U_util = rho^booked = booked service-day slots / available measured slots")
    doc.add_paragraph(
        "The guardrail is not part of the score; it determines whether the selected Q is operationally acceptable."
    )
    doc.add_picture(str(figure_paths["study_a_heatmap"]), width=Inches(6.6))
    doc.add_paragraph(
        "Figure 1. Study A objective: strict-minus-FCFS booked-slot utilization. Black x marks cells where the strict-reservation run violates the served-rate threshold."
    )

    doc.add_heading("4. Study B: Wait-Adjusted Utility With Guardrails", level=1)
    doc.add_paragraph("Main objective used to choose Q:")
    doc.add_paragraph(
        "U_wait_util = U_weighted_booked(w1,w2) - c Q/S - gamma tau_bar_w/H, with w1=1.5, w2=1.0, c=0.05, gamma=0.05."
    )
    doc.add_paragraph(
        "This score rewards booked capacity, gives Class 1 extra weight, penalizes protected capacity, and penalizes longer offered waits."
    )
    doc.add_picture(str(figure_paths["study_b_heatmap"]), width=Inches(6.6))
    doc.add_paragraph(
        "Figure 2. Study B objective: strict-minus-FCFS wait-adjusted booked-slot utility. Black x marks cells where the strict-reservation run violates the served-rate threshold."
    )

    doc.add_heading("5. Preliminary Results And Best-Q Comparison", level=1)
    doc.add_paragraph("Baseline policy comparison:")
    baseline_doc = baseline_table.copy()
    add_table(
        doc,
        baseline_doc,
        [
            "policy",
            "Q",
            "booked utilization",
            "attended utilization",
            "Class 1 served rate",
            "Class 2 served rate",
            "minimum class served rate",
            "average offered delay",
            "Study A objective",
            "Study B objective",
            "feasible under served-rate threshold?",
        ],
        digits=3,
    )
    doc.add_paragraph("Best-Q comparison by study:")
    compact_best = best_table[
        [
            "scenario",
            "best Q under Study A",
            "Study A feasible?",
            "best Q under Study B",
            "Study B feasible?",
            "do studies agree?",
            "main reason if they differ",
        ]
    ].copy()
    add_table(doc, compact_best, compact_best.columns.tolist(), digits=3)
    doc.add_picture(str(figure_paths["service_rate"]), width=Inches(6.2))
    doc.add_paragraph(
        "Figure 3. Service-rate diagnostic at the baseline scenario. Dashed lines are FCFS benchmarks; the dotted line is the 0.55 threshold."
    )
    doc.add_picture(str(figure_paths["best_q"]), width=Inches(6.6))
    doc.add_paragraph(
        "Figure 4. Best Q under each study. Black X markers with a white outline mark infeasible selected Q values."
    )
    doc.add_picture(str(figure_paths["selected_q_class_service"]), width=Inches(6.6))
    doc.add_paragraph(
        "Figure 5. Class-specific served rates at the Study B selected Q. Class 1 and Class 2 are shown on the same axes; black X markers with a white outline mark infeasible selected Q values."
    )
    doc.add_paragraph(
        f"Summary comparison: the two studies agree in {objectives_agree_count} of {objective_rows} scenarios. Study A is the more operationally defensible first-pass rule because it directly tests slot use and does not reward Class 1 access gains that come with large Class 2 service-rate losses. Study B is useful for studying priority and wait tradeoffs, but only after the service-rate guardrail is enforced."
    )

    doc.add_heading("6. Sensitivity To w1, c, gamma, And served_rate_min", level=1)
    sensitivity_doc = sensitivity[
        (sensitivity["study"] == "B: wait-adjusted")
        & (sensitivity["w1"] == HEADLINE_W1)
        & (sensitivity["c"].isin([0.0, HEADLINE_C, 0.10]))
        & (sensitivity["served_rate_min"] == HEADLINE_SERVED_RATE_MIN)
        & (sensitivity["gamma"].isin(GAMMA_VALUES))
    ][
        [
            "study",
            "w1",
            "c",
            "gamma",
            "served_rate_min",
            "best Q",
            "strict beats FCFS?",
            "binding constraint or main failure mode",
        ]
    ].copy()
    add_table(doc, sensitivity_doc, sensitivity_doc.columns.tolist(), digits=3)
    doc.add_picture(str(figure_paths["gamma_sensitivity"]), width=Inches(6.4))
    doc.add_paragraph(
        "Figure 6. Study B sensitivity at the baseline scenario. The plotted objective is the best strict-minus-FCFS wait-adjusted utility for each gamma and slot-cost value."
    )

    doc.add_heading("7. Decision Points Before Finalizing", level=1)
    add_bullets(
        doc,
        [
            "Confirm whether the service-rate floor should be 0.55 or higher.",
            "Decide whether booked utilization is sufficient, or whether attended utilization should be a hard constraint because no-shows do not produce visits.",
            "Decide whether Class 2 service-rate loss should have its own explicit threshold.",
            "Rerun with more seeds before treating any Q recommendation as final.",
        ],
    )

    doc.add_heading("Audit Notes", level=1)
    add_bullets(
        doc,
        [
            f"Existing DOCX old-objective terms found: {audit.old_doc_terms_found}.",
            f"Old net-priority plot files found and not reused: {audit.old_plot_terms_found}.",
            f"booked_slot_utilization available in metrics layer/output: {audit.booked_metric_available}.",
            f"Class offered-delay totals available: {audit.offered_delay_available}.",
            f"Class service-rate guardrails computable: {audit.service_rates_available}.",
            "Each study uses exactly one main objective for best-Q selection.",
        ],
    )

    doc.save(DOCX_PATH)

    return {
        "study_a_best_q": study_a_best_q,
        "study_b_best_q": study_b_best_q,
        "wait_changed_best_q": wait_changed_best_q,
        "study_a_delta": study_a_delta,
        "study_b_delta": study_b_delta,
        "infeasible_util_wins": infeasible_util_wins,
        "objectives_agree_count": objectives_agree_count,
        "objective_rows": objective_rows,
    }


def try_make_pdf() -> bool:
    pandoc = subprocess.run(["command", "-v", "pandoc"], shell=True, capture_output=True, text=True)
    if pandoc.returncode != 0:
        return False
    # Pandoc DOCX-to-PDF depends on a local PDF engine. Treat failure as non-blocking.
    result = subprocess.run(
        ["pandoc", str(DOCX_PATH), "-o", str(PDF_PATH)],
        cwd=REPO_DIR,
        capture_output=True,
        text=True,
    )
    return result.returncode == 0 and PDF_PATH.exists()


def main() -> None:
    ensure_dirs()
    run_df, params = load_run_data()
    audit = audit_inputs(run_df)
    if not audit.booked_metric_available:
        raise RuntimeError("booked_slot_utilization is not available.")
    if not audit.offered_delay_available:
        raise RuntimeError("Class offered-delay totals are not available.")
    if not audit.service_rates_available:
        raise RuntimeError("Class service-rate columns are not available.")

    seeds = sorted(int(seed) for seed in run_df["seed"].unique())
    scored = with_objectives(run_df, w1=HEADLINE_W1, c=HEADLINE_C, gamma=HEADLINE_GAMMA)
    study_a_summary = paired_summary(
        scored,
        objective_column="study_a_objective",
        served_rate_min=HEADLINE_SERVED_RATE_MIN,
    )
    study_b_summary = paired_summary(
        scored,
        objective_column="study_b_objective",
        served_rate_min=HEADLINE_SERVED_RATE_MIN,
    )
    best_a = choose_best_q(study_a_summary)
    best_b = choose_best_q(study_b_summary)

    baseline_table = baseline_policy_table(
        scored,
        best_a,
        best_b,
        served_rate_min=HEADLINE_SERVED_RATE_MIN,
    )
    best_table = best_q_comparison_table(best_a, best_b)
    sensitivity = baseline_sensitivity(scored)

    save_tables(baseline_table, best_table, sensitivity, study_a_summary, study_b_summary)
    figure_paths = generate_figures(scored, study_a_summary, study_b_summary, best_table, sensitivity)
    memo_stats = make_doc(
        baseline_table=baseline_table,
        best_table=best_table,
        sensitivity=sensitivity,
        figure_paths=figure_paths,
        audit=audit,
        seeds=seeds,
        params=params,
    )
    pdf_created = try_make_pdf()

    strict_beats_baseline_a = memo_stats["study_a_delta"] > EPS
    strict_beats_baseline_b = memo_stats["study_b_delta"] > EPS
    bottom_line = (
        "preliminary baseline results favor FCFS/no reservation"
        if not strict_beats_baseline_a and not strict_beats_baseline_b
        else "preliminary baseline results show a strict-reservation advantage under at least one study"
    )
    limitations = "5 seeds only; cached equal-arrival sweep; preliminary direction-finding results"

    summary = {
        "output path": str(DOCX_PATH),
        "pdf created": pdf_created,
        "old objective functions found in existing doc": audit.old_doc_terms_found,
        "old objective plots found and replaced": audit.old_plot_terms_found,
        "objective functions analyzed": [
            "Study A: U_util = rho^booked",
            "Study B: U_wait_util = U_weighted_booked - c*Q/S - gamma*tau_bar_w/H",
        ],
        "seed set used": f"{seeds[0]}-{seeds[-1]} ({len(seeds)} seeds)",
        "service-rate threshold used": HEADLINE_SERVED_RATE_MIN,
        "service-rate violations marked in plots": True,
        "wait penalty changed baseline best Q": memo_stats["wait_changed_best_q"],
        "bottom-line conclusion": bottom_line,
        "limitations": limitations,
    }
    (OUTPUT_DIR / "terminal_summary.json").write_text(json.dumps(summary, indent=2), encoding="utf-8")

    print("Terminal summary")
    for key, value in summary.items():
        print(f"- {key}: {value}")


if __name__ == "__main__":
    main()
